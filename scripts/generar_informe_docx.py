"""Genera ENTREGABLE5_INFORME_PABLOGARCIAMARQUEZ.docx (~5 paginas, evidencias clave)."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Inches, Pt

BASE = Path(__file__).resolve().parents[1]
EVID = BASE / "docs" / "evidencias"
OUT_DOCX = BASE / "docs" / "ENTREGABLE5_INFORME_PABLOGARCIAMARQUEZ.docx"
OUT_PDF = BASE / "docs" / "ENTREGABLE5_INFORME_PABLOGARCIAMARQUEZ.pdf"

IMG = Inches(4.1)
IMG_WIDE = Inches(4.6)


def style_paragraph(paragraph, bold=False):
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph.paragraph_format.space_after = Pt(4)
    for run in paragraph.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(12)
        run.bold = bold


def add_heading(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(13)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)


def add_body(doc, text):
    p = doc.add_paragraph(text)
    style_paragraph(p)


def add_figure(doc, filename, caption, width=IMG):
    path = EVID / filename
    if path.exists():
        doc.add_picture(str(path), width=width)
        cap = doc.add_paragraph(caption)
        cap.runs[0].italic = True
        style_paragraph(cap)
    else:
        add_body(doc, f"[FALTA IMAGEN: {filename}]")


def main():
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(12)

    title = doc.add_paragraph()
    t = title.add_run("ENTREGABLE 5 — Informe\n")
    t.bold = True
    t.font.size = Pt(15)
    sub = title.add_run("Agente RAG · FastAPI · Azure SQL · Docker · ACR · ACA · CI/CD")
    sub.font.size = Pt(12)
    style_paragraph(title)

    add_body(doc, "Alumno: Pablo García Márquez · Proyecto: entregable5-rag-azure")
    add_body(doc, "Fecha: 13 de julio de 2026")
    add_body(
        doc,
        "Nota: Este documento es la versión resumida (4 páginas) preparada según el "
        "requisito del enunciado (3-5 páginas). Existe una versión ampliada de 9 páginas "
        "con 17 capturas de evidencia, conservada en docs/evidencias/ del repositorio.",
    )

    add_heading(doc, "1. Objetivo y arquitectura")
    add_body(
        doc,
        "Se despliega una API RAG con FastAPI que consulta fragmentos en Azure SQL y "
        "devuelve respuestas con fuentes trazables. El flujo es: Azure Repos → pipeline "
        "DevOps (tests, build, deploy, validación) → ACR → Container App (ingress :8000) "
        "→ ragdb. Recursos en rg-entregable5-rag (France Central).",
    )
    add_figure(
        doc,
        "fase2-arquitectura-diagrama.png",
        "Figura 1. Arquitectura DevOps → ACR → ACA → Azure SQL.",
        width=IMG_WIDE,
    )

    add_heading(doc, "2. Repositorio, SQL y desarrollo local")
    add_body(
        doc,
        "Código en Azure DevOps (main): Dockerfile, docker-compose, tests y "
        "azure-pipelines.yml. Secretos en .env local y Variable Group; no en el repo. "
        "Base ragdb con tabla knowledge_chunks (4 filas). Contenerización validada con "
        "docker compose build y Swagger local (/ask con sources).",
    )
    add_figure(
        doc,
        "fase1-01-repositorio-azure-devops-codigo.png",
        "Figura 2. Repositorio y código en Azure DevOps.",
    )
    add_figure(
        doc,
        "fase2-07-sql-query-editor-chunks.png",
        "Figura 3. Datos de conocimiento en Azure SQL.",
    )
    add_figure(
        doc,
        "fase4-04-docker-compose-build.png",
        "Figura 4. docker compose build correcto.",
    )

    add_heading(doc, "3. Azure: ACR, Container Apps y pipeline")
    add_body(
        doc,
        "ACR acragent5pgar (Basic, admin user) almacena rag-agent-api. El pipeline "
        "publica tags BuildId y latest. ACA cae-entregable5-rag usa imagen de ACR, "
        "ingress externo y secretref para la cadena SQL. Agente self-hosted NAVE-NODRIZA "
        "(pool Default); run #20260713.3 con cuatro stages en verde.",
    )
    add_figure(
        doc,
        "fase5-04-acr-tag-pipeline.png",
        "Figura 5. Imagen en ACR con tag del pipeline.",
    )
    add_figure(
        doc,
        "fase6-01-aca-running-fqdn.png",
        "Figura 6. Container App en ejecución con URL pública.",
    )
    add_figure(
        doc,
        "fase7-01-pipeline-stages-verde.png",
        "Figura 7. Pipeline CI/CD completado.",
    )

    add_heading(doc, "4. Validación en cloud y reflexión")
    add_body(
        doc,
        "URL pública operativa. GET /health: database connected, 4 chunks. POST /ask: "
        "answer y sources. Logs ACA: Uvicorn y peticiones HTTP 200. "
        "RAG aporta trazabilidad; secretos con secretref y Variable Group. "
        "Azure for Students exige agente self-hosted.",
    )
    add_figure(
        doc,
        "fase8-02-ask-cloud-sources.png",
        "Figura 8. POST /ask en cloud con sources.",
    )

    add_heading(doc, "5. Conclusión")
    add_body(
        doc,
        "Ciclo completo verificado: código, Docker, Azure SQL, ACR, ACA y CI/CD con "
        "validación automática y manual en la nube. La solución integra un flujo RAG "
        "trazable automatizado hasta Azure Container Apps.",
    )

    doc.save(OUT_DOCX)
    print(f"Generado (version ~5 paginas): {OUT_DOCX}")


def export_pdf():
    import win32com.client  # type: ignore

    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    doc = word.Documents.Open(str(OUT_DOCX.resolve()))
    doc.ExportAsFixedFormat(str(OUT_PDF.resolve()), ExportFormat=17)
    doc.Close(False)
    word.Quit()
    print(f"PDF exportado: {OUT_PDF}")


if __name__ == "__main__":
    main()
